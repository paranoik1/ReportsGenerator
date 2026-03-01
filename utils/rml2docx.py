import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx_utils import setup_document_styles


RML = str


def set_paragraph_indent(paragraph, level):
    """Set left indent and hanging indent for list paragraphs."""
    indent = level * 0.25  # inches per level
    paragraph.paragraph_format.left_indent = Inches(indent)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)  # hanging indent


def add_bullet_paragraph(doc, text, level):
    """Add a bullet list paragraph with appropriate indentation."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    # Use a bullet character (-) or rely on Word's auto-formatting by setting style
    # We'll manually set bullet symbol and hanging indent for simplicity.
    run.text = f"- {text}"
    set_paragraph_indent(p, level)
    return p


def add_numbered_paragraph(doc, text, level, prefix):
    """Add a numbered list paragraph with manual prefix and hanging indent."""
    p = doc.add_paragraph()
    p.add_run(f"{prefix} {text}")
    set_paragraph_indent(p, level)
    return p


def process_text_for_bold(text, bold_active):
    """
    Process a line of text for [$ start bold $] and [$ end bold $] tags.
    Returns a list of (text_segment, is_bold) and the final bold state.
    """
    # Regex to match bold tags
    pattern = r"(\[$ start bold \$\]|\[$ end bold \$\])"
    parts = re.split(pattern, text)
    runs = []
    current_bold = bold_active
    for part in parts:
        if part == "[$ start bold $]":
            current_bold = True
        elif part == "[$ end bold $]":
            current_bold = False
        elif part:  # text segment
            runs.append((part, current_bold))
    return runs, current_bold



def rml_to_docx(rml_text: RML, output_path: str, image_base_path: str = "."):
    """
    Convert RML text to a .docx file.

    :param rml_text: string containing RML markup
    :param output_path: path where the .docx file will be saved
    :param image_base_path: base directory for image file paths (default current dir)
    """
    doc = Document()
    setup_document_styles(doc)
    lines = rml_text.splitlines()

    # State variables
    bold_active = False  # whether we are inside an unclosed bold block
    numbered_counters: dict[int, int] = (
        {}
    )  # e.g., {1: 1, 2: 2} for level -> current number

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.lstrip()

        # --- Alignment tag ---
        align_match = re.match(
            r"^\[\$ align (left|right|center|justify) \$](.*)$", line
        )
        if align_match:
            align_val = align_match.group(1)
            rest_text = align_match.group(2)
            p = doc.add_paragraph()
            # Set alignment
            if align_val == "left":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align_val == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align_val == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align_val == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Process remaining text (if any) for bold
            if rest_text.strip():
                runs, bold_active = process_text_for_bold(rest_text, bold_active)
                for text_seg, is_bold in runs:
                    run = p.add_run(text_seg)
                    run.bold = is_bold
            i += 1
            continue

        # --- Image tag ---
        img_match = re.match(r"^\[\$ img (.*?) \$]$", line)
        if img_match:
            filepath = img_match.group(1).strip()
            full_path = os.path.join(image_base_path, filepath)
            p = doc.add_paragraph()
            run = p.add_run()
            try:
                run.add_picture(full_path, width=Inches(6))  # scale to 6 inches wide
            except Exception as e:
                # If image not found, insert placeholder text
                run.add_text(f"[Image not found: {filepath}]")
            i += 1
            continue

        # --- Numbered list tag ---
        num_match = re.match(r"^\[\$ num (\d+) \$](.*)$", line)
        if num_match:
            level = int(num_match.group(1))
            rest_text = num_match.group(2).lstrip()

            # Update counters
            # Increment current level
            numbered_counters[level] = numbered_counters.get(level, 0) + 1
            # Reset any higher levels
            for lvl in list(numbered_counters.keys()):
                if lvl > level:
                    del numbered_counters[lvl]

            # Build prefix (e.g., "1.", "1.1.")
            prefix_parts = []
            for lvl in range(1, level + 1):
                if lvl in numbered_counters:
                    prefix_parts.append(str(numbered_counters[lvl]))
                else:
                    prefix_parts.append("1")  # should not happen
            prefix = ".".join(prefix_parts) + "."

            # Create paragraph with manual numbering
            p = doc.add_paragraph()
            # Process rest_text for bold
            runs, bold_active = process_text_for_bold(rest_text, bold_active)
            full_text = prefix + " "
            # Build run by run to apply bold correctly
            p.add_run(full_text)  # prefix not bold by default; can be changed if needed
            for text_seg, is_bold in runs:
                run = p.add_run(text_seg)
                run.bold = is_bold
            set_paragraph_indent(p, level)
            i += 1
            continue

        # --- Bullet list ---
        # Detect bullet lines: optional leading spaces, then hyphen
        bullet_match = re.match(r"^(\s*)- (.*)$", line)
        if bullet_match:
            spaces = bullet_match.group(1)
            rest_text = bullet_match.group(2)
            # Determine level by number of spaces (e.g., 2 spaces per level)
            level = len(spaces) // 2  # assuming 2 spaces per level, or 0 if none
            # Create bullet paragraph with manual bullet
            p = doc.add_paragraph()
            # Process rest_text for bold
            runs, bold_active = process_text_for_bold(rest_text, bold_active)
            # Add bullet symbol
            p.add_run("- ")
            for text_seg, is_bold in runs:
                run = p.add_run(text_seg)
                run.bold = is_bold
            set_paragraph_indent(p, level)
            i += 1
            continue

        # --- Normal paragraph (including empty lines) ---
        # Handle empty line
        if line == "":
            doc.add_paragraph()
            i += 1
            continue

        # Process normal text with possible bold tags
        p = doc.add_paragraph()
        runs, bold_active = process_text_for_bold(line, bold_active)
        for text_seg, is_bold in runs:
            run = p.add_run(text_seg)
            run.bold = is_bold
        i += 1

    # Save document
    doc.save(output_path)


# Example usage
if __name__ == "__main__":
    sample_rml = """[$ align center $][$ start bold $]Практическая работа №1[$ end bold $]
[$ align center $][$ start bold $]Изучение сетевых адаптеров и их характеристик[$ end bold $]

[$ start bold $]Цель работы:[$ end bold $] изучить конструктивные особенности сетевых карт, определить типы интерфейсов подключения и физических сред передачи данных, а также освоить методику просмотра параметров сетевого адаптера в операционной системе Windows, включая определение MAC-адреса.

[$ start bold $]Ход работы:[$ end bold $]

[$ num 1 $] Проведён визуальный осмотр сетевой карты, извлечённой из персонального компьютера. Определены следующие характеристики:
- Тип шины (интерфейс подключения): По длине контактной части платы (менее 10 см) установлено, что сетевая карта подключается к шине PCI.
- Тип физической среды: По типу разъёма на металлической задней панели определено, что карта предназначена для работы с витой парой (разъём RJ-45).

[$ num 1 $] Дополнительно отмечены поддерживаемые функции сетевой карты:
- Поддержка Boot ROM (загрузка ПК по сети без жёсткого диска)
- Поддержка Wake On LAN (включение ПК по сигналу из сети)
- Поддержка режима Full Duplex (одновременная передача и приём данных)
- Наличие индикаторов активности на задней панели

[$ num 1 $] Изучили сетевой адаптер в операционной системе Windows. В операционной системе Windows 10 выполнены следующие действия:
- Открыт Диспетчер устройств через меню: Пуск → Панель управления → Оборудование и звук → Диспетчер устройств
- Раскрыт раздел «Сетевые адаптеры». Установлено, что в системе обнаружено два сетевых адаптера. Отсутствие жёлтых восклицательных знаков и красных крестиков свидетельствует о корректной установке драйверов.

[$ img media/image1.png $]
[$ align center $]Рисунок 1. Сетевые адаптеры в Диспетчере устройств

[$ num 1 $] Определили MAC-адреса сетевого адаптера. В операционной системе Windows 10 выполнены следующие действия:
- Открыта командная строка: Пуск → Все программы → Стандартные → Командная строка
- Выполнена команда: ipconfig /all
- В полученном выводе найден параметр «Физический адрес», который представляет собой уникальный MAC-адрес сетевого адаптера

[$ img media/image2.png $]
[$ align center $]Рисунок 2. MAC-адрес

[$ start bold $]Вывод:[$ end bold $] в ходе практической работы изучены основные типы сетевого оборудования, в частности сетевые адаптеры. Путём визуального осмотра определены тип интерфейса подключения (PCI) и тип поддерживаемой физической среды передачи данных (витая пара, разъём RJ-45). В операционной системе Windows с помощью Диспетчера устройств подтверждена корректная работа сетевых адаптеров, а с помощью команды ipconfig /all определён уникальный MAC-адрес сетевого интерфейса. Практическая работа позволила закрепить теоретические знания о характеристиках сетевых карт и освоить базовые методы диагностики сетевого оборудования в ОС Windows.
"""

    # Save to a file (make sure the media folder and images exist, or handle missing)
    rml_to_docx(sample_rml, "output.docx", image_base_path=".")
