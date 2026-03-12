from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

from utils.docx_styles import setup_document_styles


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


class HTMLToDocx:
    def __init__(self, document: DocumentObject):
        self.doc = document

    def render(self, html: str):
        soup = BeautifulSoup(html, "html.parser")

        for element in soup.contents:
            self.handle_block(element)

    # ---------- BLOCK LEVEL ----------

    def handle_block(self, node):
        if isinstance(node, NavigableString):
            if node.strip():
                self.doc.add_paragraph(node.strip())
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        if name == "p":
            self.handle_paragraph(node)

        elif name == "center":
            self.handle_center(node)

        elif name == "ol":
            self.handle_list(node, ordered=True, level=0)

        elif name == "ul":
            self.handle_list(node, ordered=False, level=0)

        elif name == "img":
            self.handle_image(node)

        elif name == "table":
            self.handle_table(node)

        elif name == "code":
            self.handle_code_block(node)

        else:
            for child in node.children:
                self.handle_block(child)

    def handle_paragraph(self, tag: Tag):
        p = self.doc.add_paragraph()
        self.render_inline(tag, p)

    def handle_center(self, tag: Tag):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.render_inline(tag, p)

    def handle_image(self, tag: Tag):
        src = tag.get("src")
        if src:
            try:
                self.doc.add_picture(src)  # type: ignore
            except FileNotFoundError:
                print("Файл изображения не найден:", src)

    def handle_code_block(self, tag: Tag):
        """
        Обрабатывает блок кода <code>.
        Использует стиль 'Code Style', определенный в docx_styles.py.
        """
        code_text = tag.get_text()
        self.doc.add_paragraph(code_text, style="Code Style")

    def handle_list(self, tag: Tag, ordered: bool, level: int):
        style = "List Number" if ordered else "List Bullet"

        for li in tag.find_all("li", recursive=False):
            p = self.doc.add_paragraph(style=style)
            self.render_inline(li, p)

            for child in li.children:
                if isinstance(child, Tag) and child.name in ("ul", "ol"):
                    self.handle_list(child, child.name == "ol", level + 1)

    # ---------- TABLE SUPPORT ----------

    def handle_table(self, table_tag: Tag):
        rows = table_tag.find_all("tr")

        if not rows:
            return

        max_cols = 0
        table_data = []

        for tr in rows:
            cells = tr.find_all(["td", "th"], recursive=False)
            row_data = []
            for cell_tag in cells:
                row_data.append(cell_tag)
            max_cols = max(max_cols, len(row_data))
            table_data.append(row_data)

        table = self.doc.add_table(
            rows=len(table_data), cols=max_cols, style="Normal Table"
        )

        for i, row in enumerate(table_data):
            for j, cell_tag in enumerate(row):
                cell = table.cell(i, j)
                cell_style = {"sz": 5, "val": "single"}
                set_cell_border(
                    cell,
                    top=cell_style,
                    bottom=cell_style,
                    start=cell_style,
                    end=cell_style,
                )

                paragraph = cell.paragraphs[0]
                paragraph.clear()
                paragraph.style = "Table Content"

                if cell_tag.name.lower() == "th":
                    run = paragraph.add_run(cell_tag.get_text(strip=True))
                    run.bold = True
                else:
                    self.render_inline(cell_tag, paragraph)

    # ---------- INLINE ----------

    def render_inline(self, tag: Tag, paragraph):
        for child in tag.children:

            if isinstance(child, NavigableString):
                if child.strip():
                    paragraph.add_run(str(child))

            elif isinstance(child, Tag):
                self.handle_inline_tag(child, paragraph)

    def handle_inline_tag(self, tag: Tag, paragraph):
        name = tag.name.lower()

        if name in ("strong", "b"):
            run = paragraph.add_run(tag.get_text())
            run.bold = True

        elif name in ("em", "i"):
            run = paragraph.add_run(tag.get_text())
            run.italic = True

        elif name == "img":
            self.handle_image(tag)

        elif name == "center":
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.render_inline(tag, p)

        else:
            self.render_inline(tag, paragraph)

def html_to_docx(html_path: str, docx_path: str):
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    doc = Document()
    setup_document_styles(doc)
    renderer = HTMLToDocx(doc)
    renderer.render(html)

    doc.save(docx_path)


if __name__ == "__main__":
    import markdown

    with open("tmp/report2.md") as fp:
        md_result = fp.read()

    html_result = markdown.markdown(
        md_result, extensions=["extra", "sane_lists", "nl2br"]
    )

    with open("tmp/report2.html", "w") as fp:
        fp.write(html_result)

    html_to_docx("tmp/report2.html", "tmp/report2.docx")
